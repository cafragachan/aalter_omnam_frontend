"""
Generates the engineering handoff report as a .docx file at the project root.

    python _build_handoff_report.py

Re-run after editing content below to regenerate.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT_PATH = Path(__file__).parent / "Omnam Frontend - Engineering Handoff.docx"

# --------------------------------------------------------------------------- #
# Palette + styling helpers (adapted from reports/_build_voice_latency_report) #
# --------------------------------------------------------------------------- #
ACCENT = RGBColor(0x1F, 0x3A, 0x68)   # dark blue
ACCENT2 = RGBColor(0x2E, 0x6B, 0x5E)  # teal (sub-accents)
MUTED = RGBColor(0x55, 0x55, 0x55)
CODECLR = RGBColor(0x9B, 0x26, 0x15)  # code/inline literal color
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_table(doc, headers, rows, col_widths=None, header_fill="1F3A68", font=9.5):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    table.autofit = True
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(font)
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r % 2 == 0:
                shade_cell(cell, "F2F5F9")
    doc.add_paragraph()
    return table


def add_para(doc, text, bold=False, italic=False, size=10.5, color=None, after=6, before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    return p


def add_rich(doc, segments, size=10.5, after=6, style=None):
    """segments = list of (text, kind) where kind in {'b','i','c',''}."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for text, kind in segments:
        run = p.add_run(text)
        run.font.size = Pt(size)
        if kind == "b":
            run.bold = True
        elif kind == "i":
            run.italic = True
        elif kind == "c":
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODECLR
    p.paragraph_format.space_after = Pt(after)
    return p


def add_h1(doc, text, num=None):
    p = doc.add_paragraph()
    run = p.add_run((f"{num}.  " if num else "") + text)
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    add_hr(doc)


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(3)


def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT2
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def add_bullet(doc, text=None, segments=None, size=10.5, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    if segments:
        return add_rich(doc, segments, size=size, after=2, style=style)
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F4F4F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    return p


# --------------------------------------------------------------------------- #
# Document                                                                     #
# --------------------------------------------------------------------------- #
doc = Document()
section = doc.sections[0]
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ---- Title block ----------------------------------------------------------- #
title = doc.add_paragraph()
trun = title.add_run("Omnam Metaverse — Frontend Engineering Handoff")
trun.bold = True
trun.font.size = Pt(23)
trun.font.color.rgb = ACCENT

sub = doc.add_paragraph()
srun = sub.add_run("Architecture, AI Pipeline & Latency Reference for the Incoming Engineer")
srun.font.size = Pt(12.5)
srun.font.color.rgb = MUTED
srun.italic = True

meta = doc.add_paragraph()
mrun = meta.add_run(
    "Repository: aalter_omnam_frontend   ·   Branch: main   ·   Version: 0.4.3   ·   Prepared: June 2026"
)
mrun.font.size = Pt(9.5)
mrun.font.color.rgb = MUTED

add_para(doc,
         "Mandate for the incoming engineer: (1) reduce conversational latency, (2) improve the AI "
         "agent's behaviour, and (3) improve guest data extraction and analysis. This document maps "
         "the whole frontend, the UE5 and AI integrations, and points directly at the code paths and "
         "trade-offs that bear on those three goals.",
         italic=True, color=MUTED, before=4)
add_hr(doc)

# =========================================================================== #
# 1. EXECUTIVE SUMMARY
# =========================================================================== #
add_h1(doc, "Executive Summary", 1)
add_para(doc,
    "Omnam is a Next.js 16 / React 19 web app that lets a guest book a luxury hotel stay by talking "
    "to an AI concierge avatar (\"Ava\") while a photoreal Unreal Engine 5 digital twin of the hotel "
    "is streamed into the page. The guest speaks; HeyGen LiveAvatar transcribes and lip-syncs the "
    "avatar; a server-side orchestration LLM decides what Ava says and which 3D scene to navigate "
    "to; commands are pushed to UE5 over a message channel; and the whole journey is persisted to "
    "Firebase for returning-guest personalization and operator analytics.")
add_para(doc,
    "The system is built around a pure state machine (the \"journey reducer\") that is fed by a "
    "hybrid intent layer: fast deterministic regex paths handle roughly half of turns in well under "
    "a second, and an LLM handles the ambiguous remainder. The dominant latency cost is the LLM "
    "round-trip plus HeyGen's ~700 ms voice-synthesis warmup. No LLM endpoint currently streams, "
    "which is the single largest available latency win.")

add_h2(doc, "The five external systems you must hold in your head")
add_bullet(doc, segments=[("HeyGen LiveAvatar ", "b"),
    ("(", ""), ("@heygen/liveavatar-web-sdk", "c"), (") — speech-to-text, the avatar's voice (TTS), "
    "and the green-screen video that is chroma-keyed onto a canvas in the browser.", "")])
add_bullet(doc, segments=[("Unreal Engine 5 + Vagon.io ", "b"),
    ("— the 3D hotel, pixel-streamed into an ", ""), ("<iframe>", "c"),
    (". In production the stream is hosted on Vagon; in local dev the app talks to a UE5 editor over ", ""),
    ("ws://localhost:7788", "c"), (".", "")])
add_bullet(doc, segments=[("OpenAI + Anthropic ", "b"),
    ("— two LLM providers used across nine API routes (orchestration, room planning, intent, speech, "
     "guest analysis, points-of-interest).", "")])
add_bullet(doc, segments=[("Google Places API ", "b"),
    ("— verifies and geo-locates nearby points of interest for the ", ""),
    ("location", "c"), (" experience.", "")])
add_bullet(doc, segments=[("Firebase (Auth + Realtime DB + Storage) ", "b"),
    ("— guest identity, learned personality/preferences, and full per-session transcripts/snapshots.", "")])

add_para(doc,
    "Important: the checked-in CLAUDE.md is partly out of date. It describes a smaller system "
    "(\"gpt-5.4-mini via /api/orchestrate\", an EventBus, only /api/orchestrate + "
    "/api/start-sandbox-session). The live code has nine API routes, a Firebase persistence layer, "
    "Vagon lifecycle management, guest-intelligence analytics, and a different model line-up. This "
    "report reflects the code as it is on disk.", italic=True, color=MUTED, before=4)

# =========================================================================== #
# 2. TECH STACK
# =========================================================================== #
add_h1(doc, "Technology Stack", 2)
add_table(doc,
    ["Layer", "Choice", "Notes"],
    [
        ["Framework", "Next.js 16 (App Router), React 19, TypeScript 5", "Server routes under app/api/* run on the Next server."],
        ["Styling", "Tailwind CSS v4, glassmorphism design system", "shadcn/ui (~50 components) in components/ui/."],
        ["AI avatar", "HeyGen LiveAvatar Web SDK ^0.0.10", "Voice chat, TTS, chroma-key canvas render."],
        ["3D / streaming", "Unreal Engine 5 via Vagon.io pixel stream", "iframe + Vagon SDK; local dev over WebSocket."],
        ["LLMs", "OpenAI (gpt-5.4 family, gpt-4o) + Anthropic (claude-haiku-4-5)", "Per-route; see AI Architecture section."],
        ["State", "Custom reducer store (OmnamStore) + React context", "Redux-like; synchronous effect collection."],
        ["Persistence", "Firebase Auth + Realtime DB + Cloud Storage", "Incremental writes during session + end snapshot."],
        ["Maps/POI", "Google Places API (searchText)", "Used by /api/locate-interest-points."],
        ["Analytics", "@vercel/analytics; custom latency logging", "Per-session latency logs to disk / Vercel Blob."],
    ],
    col_widths=[Cm(3.2), Cm(6.5), Cm(7.5)])
add_rich(doc, [("Note on legacy code: ", "b"),
    ("the repo still contains an abandoned migration to LiveKit + Hedra + OpenAI-Realtime "
     "(the ", ""), ("@livekit/*", "c"), (" and ", ""), ("@livekit/agents-plugin-hedra", "c"),
    (" dependencies, plus an ", ""), ("npm run agent", "c"),
    (" script). That path is dead — only the HeyGen ", ""), ("/home", "c"),
    (" route is live. Vagon HMAC-API files (", ""), ("lib/vagon-api.ts", "c"), (", ", ""),
    ("lib/vagon-client.ts", "c"), (", ", ""), ("lib/ue5/useVagonSession.ts", "c"),
    (", ", ""), ("app/api/vagon-init", "c"), (", ", ""), ("app/api/stop-vagon-machine", "c"),
    (") are marked @deprecated 2026-05-15 and unused; production uses a hosted Vagon stream URL with ", ""),
    ("?newSession=true", "c"), (".", "")], after=4)

# =========================================================================== #
# 3. HIGH-LEVEL ARCHITECTURE
# =========================================================================== #
add_h1(doc, "High-Level Architecture & Integration Map", 3)
add_para(doc, "The runtime is a single page (/home) coordinating five external systems through a "
    "central orchestration hook. The diagram below shows the data flow for one conversational turn.")

add_code_block(doc,
"""                         ┌──────────────────────────────────────────────┐
                         │              BROWSER  (/home)                 │
                         │                                              │
   guest voice  ───────► │  HeyGen SDK ──STT──► messages[] ──┐          │
                         │  (LiveAvatarContext)              │          │
                         │                                   ▼          │
   avatar video ◄─chroma─│  <canvas>            ┌────── useJourney() ───┐│
   + TTS audio  ◄────────│  (SandboxLiveAvatar) │  classifyIntent()     ││
                         │                      │  deterministic gates  ││
                         │  UE5 <iframe> ◄──cmds┤  journeyReducer (pure)││
   3D pixels    ────────►│  (Vagon stream)      │  speech mutex         ││
                         │         ▲            └───────┬───────────────┘│
                         └─────────│────────────────────│────────────────┘
                                   │ WebSocket /         │ HTTPS (fetch)
                                   │ Vagon SDK msgs      ▼
                    ┌──────────────┴───┐   ┌─────────────────────────────────────┐
                    │   UNREAL ENGINE 5 │   │     NEXT.JS API ROUTES (app/api)     │
                    │  (Vagon machine)  │   │  orchestrate · room-planner ·        │
                    └───────────────────┘   │  classify-intent · generate-speech · │
                                            │  locate-interest-points · analyze-   │
                                            │  guest · start-sandbox-session · ... │
                                            └───────┬─────────────┬──────────┬─────┘
                                                    ▼             ▼          ▼
                                              OpenAI / Anthropic  Google    HeyGen
                                                                  Places    token API
                                                    │
                                                    ▼
                                      Firebase  (Auth · Realtime DB · Storage)
                                      identity · personality · sessions · transcripts""")

add_h2(doc, "Integration responsibilities at a glance")
add_table(doc,
    ["Integration", "Direction", "Transport", "Carries"],
    [
        ["HeyGen LiveAvatar", "bidirectional", "WebRTC + SDK events", "Guest STT in; avatar video + TTS out; speak/talk state events."],
        ["UE5 / Vagon", "bidirectional", "Vagon SDK msgs (prod) / ws://localhost:7788 (dev)", "Scene-navigation commands out; 'unit selected' events in."],
        ["OpenAI", "request/response", "HTTPS fetch", "Orchestration (non-profile), room plan, intent, speech, POI, guest analysis."],
        ["Anthropic", "request/response", "HTTPS fetch", "Orchestration during PROFILE_COLLECTION (claude-haiku-4-5)."],
        ["Google Places", "request/response", "HTTPS fetch", "Verify + geolocate nearby points of interest."],
        ["Firebase", "read/write", "Firebase SDK", "Auth, returning-guest profile hydration, session persistence, operator analytics."],
    ],
    col_widths=[Cm(3.3), Cm(2.7), Cm(5.0), Cm(6.2)])

# =========================================================================== #
# 4. TURN LIFECYCLE
# =========================================================================== #
add_h1(doc, "The Conversational Turn Lifecycle", 4)
add_para(doc, "Understanding this sequence is the key to everything else — latency, AI behaviour, and "
    "data extraction all live inside it. A single turn flows through the layers below. The hook that "
    "owns it is lib/orchestrator/useJourney.ts.")

add_table(doc,
    ["#", "Step", "Where", "Cost"],
    [
        ["1", "Guest speaks; HeyGen emits USER_SPEAK_STARTED / ENDED, then USER_TRANSCRIPTION; text lands in messages[].", "lib/liveavatar/context.tsx", "STT finalization (HeyGen VAD)"],
        ["2", "useJourney debounces rapid VAD fragments (2.5 s in PROFILE_COLLECTION; unified debounce elsewhere) and mints a turn id.", "useJourney.ts", "debounce window"],
        ["3", "classifyIntent() — pure regex — returns a typed UserIntent.", "lib/orchestrator/intents.ts", "<1 ms"],
        ["4", "Deterministic gate decides: handle locally or escalate to the LLM (profile fast-path / exploration-deterministic).", "profileFastPath.ts, explorationDeterministic.ts", "<1 ms"],
        ["5", "If escalating: POST /api/orchestrate with message, journey context, profile, transcript, returning-guest data.", "useJourney.ts → app/api/orchestrate", "network + LLM (dominant)"],
        ["6", "Server builds a stage-specific system prompt, calls the LLM with a tool schema, validates the tool call (Zod + profile validator).", "app/api/orchestrate/route.ts", "LLM inference"],
        ["7", "Result dispatched into journeyReducer → produces effects (SPEAK, UE5_COMMAND, OPEN_PANEL, ...).", "journey-machine.ts", "<10 ms"],
        ["8", "Effects executed: speech claims the per-turn mutex and calls repeat(); UE5 commands postMessage to the iframe; panels open.", "useJourney.ts, ue5/bridge.ts", "—"],
        ["9", "HeyGen synthesizes audio; AVATAR_SPEAK_STARTED fires; the latency row is logged.", "liveavatar + debug.ts", "~700 ms TTS warmup"],
    ],
    col_widths=[Cm(0.8), Cm(9.0), Cm(4.6), Cm(3.0)])

add_rich(doc, [("Two response pathways. ", "b"),
    ("Roughly half of turns are resolved deterministically (steps 1-4, 7-9, no network) and respond "
     "in ~900 ms. The other half take the LLM path (all steps) and respond in ~2.2 s typical, with a "
     "tail to 6 s. Per-turn LLM calls: 0 for deterministic, 1 for the LLM path, +1 if the room "
     "planner fires, +1 (plus Places calls) for points-of-interest discovery.", "")])

# =========================================================================== #
# 5. AI ARCHITECTURE
# =========================================================================== #
add_h1(doc, "AI Architecture (deep dive)", 5)

add_h2(doc, "5.1  Model line-up — every LLM call in the system")
add_para(doc, "All values verified against source. No endpoint streams its response today (this is the "
    "biggest latency opportunity — see §11). Each LLM route uses a 5-15 s AbortController timeout.")
add_table(doc,
    ["Route", "Provider / Model", "Temp", "Purpose", "On hot path?"],
    [
        ["/api/orchestrate (PROFILE_COLLECTION)", "Anthropic claude-haiku-4-5 (max_tokens 360, tool_choice any)", "0.2", "Extract profile fields + decide next question.", "Yes — every profile turn"],
        ["/api/orchestrate (all other stages)", "OpenAI gpt-5.4-nano (tools, tool_choice)", "0.3", "Decide Ava's speech + which navigation tool to call.", "Yes — every non-deterministic turn"],
        ["/api/room-planner", "OpenAI gpt-4o (tool propose_room_plan)", "—", "Recommend a room mix from party + transcript.", "Yes — when rooms panel open"],
        ["/api/classify-intent", "OpenAI gpt-5.4-mini (5 s timeout)", "0", "LLM intent classifier (tiebreaker / experimental).", "Off main path"],
        ["/api/classify-room-plan", "OpenAI gpt-5.4-mini", "—", "Interpret room-plan edit utterances.", "Conditional"],
        ["/api/generate-speech", "OpenAI gpt-5.4-mini", "—", "Standalone speech generation helper.", "Conditional"],
        ["/api/locate-interest-points", "OpenAI gpt-5.4-mini + Google Places searchText", "—", "Propose POIs (LLM) then verify + geolocate (Places).", "Location experience"],
        ["/api/analyze-guest", "OpenAI gpt-5.4-mini (json_object)", "0.2", "Post-session personality + travel-driver extraction.", "End of session"],
    ],
    col_widths=[Cm(4.3), Cm(5.4), Cm(1.0), Cm(4.3), Cm(2.4)])
add_rich(doc, [("Watch-out: ", "b"),
    ("a stale comment in ", ""), ("orchestrate/route.ts:2666", "c"),
    (" claims PROFILE_COLLECTION uses gpt-5.4-mini; the actual ", ""), ("isProfileCollection", "c"),
    (" branch calls Anthropic claude-haiku-4-5. The comment also documents the history: gpt-4o-mini "
     "once looped ask_next/ready with empty profileUpdates, which is why a validator backstop exists. "
     "If the ", ""), ("validatorOverride", "c"),
    (" rate in the [ORCHESTRATE] logs spikes past ~30% of profile turns, the model is the suspect.", "")],
    after=4)

add_h2(doc, "5.2  Intent classification — the deterministic front door")
add_rich(doc, [("classifyIntent() ", "b"), ("(", ""), ("lib/orchestrator/intents.ts", "c"),
    (") is a pure, regex-based function returning a typed union (ROOMS, AMENITIES, AMENITY_BY_NAME, "
     "LOCATION, INTERIOR, EXTERIOR, BACK, HOTEL_EXPLORE, BOOK, the ROOM_* distribution intents, "
     "LIGHTING_*, LOCATE_INTEREST_POINTS, TRAVEL_TO_HOTEL, RETURN_TO_LOUNGE, END_EXPERIENCE, "
     "AFFIRMATIVE, NEGATIVE, OTHER_OPTIONS, DOWNLOAD_DATA, UNKNOWN). Runs in <1 ms with no I/O.", "")])
add_bullet(doc, segments=[("Priority-ordered: ", "b"),
    ("admin/exit/lounge first, then a question-guard (if the utterance looks like a question but has "
     "no action verb, navigation intents are skipped so Ava answers instead of teleporting), then "
     "navigation, then room-distribution, then affirmative/negative, then UNKNOWN.", "")])
add_bullet(doc, segments=[("Two deterministic gates sit behind it: ", "b"),
    ("profileFastPath.ts (PROFILE_COLLECTION — fires a canned next-question instantly for 6 known "
     "fields) and explorationDeterministic.ts (HOTEL_EXPLORATION / AMENITY_VIEWING — handles "
     "high-confidence navigation without the LLM). Both emit a confidence score; low confidence or "
     "UNKNOWN escalates to /api/orchestrate.", "")])
add_bullet(doc, segments=[("Designed to be swapped: ", "b"),
    ("there is also an LLM classifier at /api/classify-intent (gpt-5.4-mini) intended as a tiebreaker. "
     "It is largely off the hot path; the client sends its regex guess as a ", ""),
    ("regexHint", "c"), (" to orchestrate instead.", "")])

add_h2(doc, "5.3  The journey state machine (pure reducer)")
add_rich(doc, [("journeyReducer(state, action) → { nextState, effects[] } ", "b"),
    ("in ", ""), ("lib/orchestrator/journey-machine.ts", "c"),
    (" is a pure function with no React and no I/O — the heart of the orchestration and the most "
     "testable part of the system. Effects are descriptive objects (SPEAK_INTENT, UE5_COMMAND, "
     "OPEN_PANEL, CLOSE_PANELS, FADE_TRANSITION, SELECT_HOTEL, FETCH_POIS, ...) that useJourney "
     "interprets.", "")])
add_para(doc, "Journey stages (the conversation's spine):", before=2, after=2, bold=True, size=10.5)
add_code_block(doc,
"PROFILE_COLLECTION ─► DESTINATION_SELECT ─► VIRTUAL_LOUNGE ─► HOTEL_EXPLORATION ─┬─► ROOM_SELECTED\n"
"  (dates, guests,      (hotel grid)         (asking /         (rooms/amenities/  └─► AMENITY_VIEWING\n"
"   composition,                              exploring)        location panels,    ─► (LOUNGE/END_\n"
"   purpose, rooms)                                             UE5 navigation)        CONFIRMING) ─► END")
add_rich(doc, [("Pre-hotel shortcut: ", "b"),
    ("if a guest asks for content (\"show me the rooms\") while still in PROFILE_COLLECTION or "
     "VIRTUAL_LOUNGE, the reducer emits a Phase-1 travel + welcome immediately, and useJourney fires "
     "a Phase-2 navigation ~1.2 s later (after a ", ""), ("UE5_POST_TRAVEL_DELAY_MS = 3500", "c"),
    (" server-travel window). This is a known source of timing fragility.", "")])

add_h2(doc, "5.4  Speech authority & the per-turn mutex")
add_para(doc, "Multiple async paths can race to make Ava talk in one turn (deterministic reducer "
    "speech, LLM profile_turn speech, LLM no_action speech, room-planner speech, fast-path canned "
    "speech). lib/orchestrator/speech-mutex.ts enforces one winner per utterance:")
add_bullet(doc, segments=[("beginUtteranceTurn() ", "c"), ("mints a monotonic turn id at the top of the user-message effect.", "")])
add_bullet(doc, segments=[("Each producer calls ", ""), ("claimSpeech(turnId, source) ", "c"),
    ("before ", ""), ("repeat()", "c"),
    ("; the first claimant wins, others are suppressed and logged [SPEECH_LOCK_SUPPRESSED]. Same "
     "source may re-claim (loading + result pairs).", "")])
add_bullet(doc, segments=[("CANONICAL_REDUCER_INTENTS ", "c"),
    ("(AMENITIES, ROOMS, LOCATION, BACK, BOOK, HOTEL_EXPLORE, OTHER_OPTIONS): for these the reducer's "
     "ground-truth speech (real amenity names, prices) always wins over the LLM's paraphrase, by "
     "nulling the pre-generated speech ref.", "")])

add_h2(doc, "5.5  Server prompt construction")
add_para(doc, "app/api/orchestrate/route.ts (~2,800 lines) assembles a stage-specific system prompt "
    "from layered blocks: persona (Ava, the warm luxury concierge), returning-guest intelligence "
    "(identity/personality/preferences/loyalty), transcript window, hotel overview, amenity catalog "
    "(navigable vs. described-only), focused-amenity injection, room catalog + selected room, "
    "guest-context, a server-side reconstructed profile (see §6.2), per-stage guidance, and the "
    "client's regexHint as a tiebreaker. lib/avatar-context-builder.ts holds the persona + "
    "instruction layers.")
add_para(doc, "Tools offered depend on stage: profile_turn (PROFILE_COLLECTION), navigate_and_speak / "
    "no_action_speak (general), and an experimental action-dispatch tool family. Tool args are "
    "validated with Zod; profile_turn results pass an additional validateProfileTurn() backstop that "
    "rejects hallucinated children's ages and premature \"ready\" decisions, and strips preambles "
    "(\"Got it\", \"Perfect\").")

add_h2(doc, "5.6  Specialist AI flows")
add_bullet(doc, segments=[("Room planner ", "b"),
    ("(useRoomPlanner.ts + /api/room-planner, gpt-4o): fires when the rooms panel opens or on a "
     "room-edit utterance; returns { plan[], speech } plus computed totalPerNight / capacityOk; "
     "dispatches SET_ROOM_PLAN. An AbortController cancels stale in-flight requests so a manual edit "
     "isn't clobbered.", "")])
add_bullet(doc, segments=[("Points of interest ", "b"),
    ("(/api/locate-interest-points): LLM proposes named places (no coordinates, to avoid "
     "hallucination), Google Places searchText verifies + geolocates each, Haversine distance from "
     "the Lake Como anchor is computed, results are sent to UE5 as an ", ""), ("osm_data", "c"),
    (" JSON payload. 2-3 serial network calls — a latency hotspot.", "")])
add_bullet(doc, segments=[("Idle re-engagement ", "b"),
    ("(idle-detection.ts + reengage-prompts.ts): per-stage idle thresholds (12-18 s) fire stage-"
     "specific canned prompts when neither party is talking.", "")])
add_bullet(doc, segments=[("Post-session analysis ", "b"),
    ("(/api/analyze-guest, gpt-5.4-mini): extracts 3-5 personality traits + a primary travel driver "
     "from the transcript and analytics at session end, feeding returning-guest personalization.", "")])

# =========================================================================== #
# 6. DATA EXTRACTION & ANALYSIS
# =========================================================================== #
add_h1(doc, "Data Extraction & Analysis Pipeline", 6)
add_para(doc, "This is one of the three explicit mandates. Extraction is a three-tier hybrid: fast "
    "client regex (advisory), server-side transcript reconstruction (reconciliation), and the LLM "
    "(authoritative for the six profile-collection fields). Analysis is split into live behavioural "
    "tracking and end-of-session LLM enrichment.")

add_h2(doc, "6.1  Tier 1 — client regex extraction (advisory)")
add_rich(doc, [("lib/liveavatar/useUserProfile.ts ", "c"),
    ("scans user messages with regex and produces an AvatarDerivedProfile: name, party size, "
     "destination, start/end dates, interests (21 keywords), dietary restrictions, accessibility "
     "needs, travel purpose (24 keywords → 5 categories), budget, nationality, room-type preference, "
     "arrival time, guest composition (adults/children/ages), and room allocation. Runs <10 ms, "
     "fully synchronous.", "")])
add_rich(doc, [("Critical contract: ", "b"),
    ("components/ProfileSync.tsx syncs only passive-observation fields (interests, destination, "
     "budget, dietary, accessibility, nationality, arrival) into global state. It NEVER writes the "
     "six profile-collection fields (startDate, endDate, partySize, guestComposition, travelPurpose, "
     "roomAllocation) — orchestrate is their sole authoritative writer. Respect this when changing "
     "extraction or you will create double-write races.", "")])

add_h2(doc, "6.2  Tier 2 — server transcript reconstruction (reconciliation)")
add_para(doc, "Because React state lags the transcript, orchestrate/route.ts rebuilds an effective "
    "profile from conversationHistory before each call: it derives effectivePartySize from "
    "guestComposition, and scans for keyword flags (mentionsParty/Dates/Children/Purpose/Rooms). "
    "These hints are rendered into the prompt with the instruction \"transcript wins on disagreement; "
    "verify and re-extract.\" This is what keeps the agent from re-asking already-answered questions.")

add_h2(doc, "6.3  Tier 3 — LLM extraction (authoritative)")
add_para(doc, "During PROFILE_COLLECTION the claude-haiku-4-5 profile_turn tool returns "
    "profileUpdates + a decision (ask_next / clarify / ready) + speech. validateProfileTurn() then "
    "enforces invariants the model cannot be trusted on: childrenAges length must equal children "
    "count, \"ready\" requires all six fields, and decision is corrected when the next field is "
    "actually already complete.")

add_h2(doc, "6.4  Behavioural analysis & guest intelligence")
add_rich(doc, [("lib/guest-intelligence.ts ", "c"),
    ("passively tracks, during the conversation: upsell receptivity, top questions, objections "
     "(topic / resolution / resolved), booking outcome, conversation duration, rooms & amenities "
     "explored with time-spent timers, consent flags, requirements, referral source, device, and "
     "(post-analysis) personality traits + travel driver. getDataSnapshot() exposes the live state, "
     "including in-progress timers.", "")])
add_para(doc, "This data drives two things: the end-of-session /api/analyze-guest enrichment, and the "
    "Firebase-persisted personality/preferences that personalize a returning guest's next visit "
    "(merged via union-sets in lib/firebase/user-profile-service.ts). The operator-facing /admin page "
    "renders all of it per guest and per session.")

# =========================================================================== #
# 7. UE5 + VAGON
# =========================================================================== #
add_h1(doc, "UE5 / Vagon Integration", 7)
add_para(doc, "The 3D twin is a pixel-streamed UE5 build. lib/ue5/bridge.ts (useUE5Bridge) is the "
    "typed command layer over lib/useUE5WebSocket.ts, which is a dual transport: the Vagon SDK "
    "(window.Vagon.sendApplicationMessage) in production, or a raw WebSocket to ws://localhost:7788 "
    "in local dev (toggled by NEXT_PUBLIC_STREAM_MODE).")

add_h2(doc, "7.1  Message protocol")
add_h3(doc, "Outgoing (Frontend → UE5)")
add_table(doc,
    ["type", "value", "Effect"],
    [
        ["startTEST", '"startTEST"', "Begin the experience / enter the hotel."],
        ["gameEstate", '"rooms" | "amenities" | "location" | "default"', "Top-level scene navigation."],
        ["selectedRoom", "comma-separated room-type ids", "Highlight the recommended room set (driven by currentRoomPlan)."],
        ["unitView", '"interior" | "exterior"', "Switch the selected unit's camera."],
        ["communal", "amenityId (or \"virtualLounge\")", "Navigate to a specific amenity / lounge scene."],
        ["sunPosition", '"daylight" | "sunset" | "night"', "Change lighting (SunToggle)."],
        ["hotelSelected", "hotel slug", "Set the active property context."],
        ["osm_data", "JSON string of POI markers", "Initialise AOSMInterestPointsManager with located POIs."],
        ["inputRelease", "(no value)", "Release held mouse input — cross-origin iframe fix."],
    ],
    col_widths=[Cm(2.8), Cm(6.5), Cm(7.0)])
add_h3(doc, "Incoming (UE5 → Frontend)")
add_table(doc,
    ["type", "Shape", "Handling"],
    [
        ["unit", "{ roomName, description?, price?, level? }", "User clicked a unit in 3D → dispatched as UNIT_SELECTED_UE5 → opens UnitDetailPanel."],
        ["(generic)", "{ type, ... }", "Passed to the onMessage callback for other UE5 telemetry."],
    ],
    col_widths=[Cm(2.8), Cm(7.0), Cm(6.5)])

add_h2(doc, "7.2  Lifecycle, fade & latency-relevant constants")
add_bullet(doc, segments=[("Production stream: ", "b"),
    ("a hosted Vagon URL with ", ""), ("?newSession=true", "c"),
    (" forces a fresh machine on every load; Vagon's own page provisions the VM (the old client-side "
     "getStreams→startMachine→assignMachine HMAC flow is deprecated/unused).", "")])
add_bullet(doc, segments=[("Vagon SDK readiness ", "b"),
    ("is polled with a 1000 ms retry loop until window.Vagon's methods exist — a fixed startup cost.", "")])
add_bullet(doc, segments=[("Fade transitions ", "b"),
    ("in bridge.ts are hardcoded: ~2000 ms to fade, 3000 ms total before the overlay is removed — "
     "applied on every scene change, so navigation can feel sluggish. Tunable.", "")])
add_bullet(doc, segments=[("UE5_POST_TRAVEL_DELAY_MS = 3500 ", "c"),
    ("(useJourney.ts) gates the pre-hotel shortcut's second phase.", "")])
add_bullet(doc, segments=[("No explicit WebSocket reconnection ", "b"),
    ("logic — the app relies on the Vagon SDK's onDisconnected callback.", "")])

# =========================================================================== #
# 8. STATE & PERSISTENCE
# =========================================================================== #
add_h1(doc, "State Management & Persistence", 8)
add_h2(doc, "8.1  OmnamStore — the single source of truth")
add_rich(doc, [("lib/omnam-store.tsx ", "c"),
    ("is a Redux-like store holding { profile, journeyStage, app, journey, currentRoomPlan }. Its "
     "dispatch is unusual and important: it runs journeyReducer and the root reducer against a live "
     "stateRef synchronously to capture effects before React re-renders, then schedules the React "
     "dispatch. This lets async callbacks (HeyGen event handlers) read fresh state and avoids "
     "StrictMode double-invocation. ", "")])
add_rich(doc, [("lib/context.tsx ", "c"), ("(UserProfile) and ", ""), ("lib/store.tsx ", "c"),
    ("(App) are thin compat shims that read from OmnamStore. currentRoomPlan has two writers — the "
     "planner (source: 'planner') and the user card controls via EDIT_ROOM_PLAN (source: 'user') — "
     "and the panel re-plan trigger skips when source is 'user' so manual edits survive reopen.", "")])

add_h2(doc, "8.2  Firebase persistence")
add_para(doc, "Two strategies exist; the live one is incremental (lib/firebase/"
    "useIncrementalPersistence.ts). Realtime DB holds identity / personality / preferences / consent "
    "/ loyalty / activeSession under omnam/users/{uid}/; full immutable session snapshots go to Cloud "
    "Storage at sessions/{userId}/{sessionId}.json with a lightweight pointer in RTDB.")
add_table(doc,
    ["Write trigger", "Frequency", "Latency / cost note"],
    [
        ["Session start (pointer + totalSessions++)", "once", "low"],
        ["Stage change / hotel select", "per transition", "medium (snapshot upload)"],
        ["Profile change → personality + preferences merge", "debounced 3 s", "HIGH — read-modify-write union per slice"],
        ["Conversation messages → activeSession", "debounced 1.8 s", "rewrites whole array (append-only would be cheaper)"],
        ["Consent / booking outcome", "on change", "low"],
        ["End of session", "once", "HIGH — /api/analyze-guest LLM + final merges"],
        ["Tab hidden (visibilitychange)", "on hide", "flushes debounces + snapshot"],
    ],
    col_widths=[Cm(6.5), Cm(3.5), Cm(6.3)])
add_para(doc, "Returning guests are hydrated on mount via useActiveSessionHydration (resumes a <30 min "
    "non-terminal activeSession) and a returning-user profile pre-fill from persisted "
    "personality/preferences.", size=10)

# =========================================================================== #
# 9. UI / PAGE FLOW
# =========================================================================== #
add_h1(doc, "UI & Page Flow", 9)
add_h2(doc, "9.1  Provider tree (layout.tsx → home)")
add_code_block(doc,
"<AuthProvider>                       app/layout.tsx — Firebase auth + identity\n"
"  <OmnamStoreProvider>               unified state (profile/app/journey/roomPlan)\n"
"    <GuestIntelligenceProvider>      behavioural tracking\n"
"      <Toaster/> <Logo/>\n"
"      ── on /home, after intro + auth + session token: ──\n"
"      <HydratedLiveAvatarContext>    HeyGen session (gated on hydration)\n"
"        <InputModeProvider>          voice vs. chat toggle\n"
"          <HomePageContent/>         layout shell: panels, avatar bar, overlays")
add_para(doc, "lib/context.tsx and lib/store.tsx are compat shims layered for backward compatibility — "
    "stacking them is a safe no-op.", size=10, color=MUTED)

add_h2(doc, "9.2  Navigation & session boot")
add_bullet(doc, segments=[("/", "c"), (" redirects to ", ""), ("/home", "c"),
    ("; ", ""), ("/login", "c"), (" also redirects to ", ""), ("/home", "c"),
    (". ", ""), ("/admin", "c"), (" is the operator analytics console; ", ""),
    ("/rooms-test", "c"), (" is an isolated panel-layout sandbox.", "")])
add_para(doc, "The /home boot sequence is the critical path to a usable avatar (~5-15 s):", before=2, after=2)
add_table(doc,
    ["Phase", "Step", "Blocking?"],
    [
        ["0", "Providers + Firebase auth ready", "yes (~500 ms)"],
        ["1", "UE5 iframe loads (Vagon / local) in background", "no (1-5 s)"],
        ["2", "Login overlay: intro video → typewriter → auth form", "until intro + auth done"],
        ["3", "First UE5 message → ue5Ready = true", "gates session fetch"],
        ["4", "POST /api/start-sandbox-session (HeyGen token + hotel catalog)", "blocks avatar (1-2 s)"],
        ["5", "useActiveSessionHydration completes", "blocks LiveAvatar provider (0.1-0.5 s)"],
        ["6", "Returning-user profile hydration", "async, non-blocking"],
        ["7", "Avatar greeting on VoiceChatState ACTIVE", "non-blocking"],
    ],
    col_widths=[Cm(1.4), Cm(11.0), Cm(3.9)])

add_h2(doc, "9.3  Panels")
add_bullet(doc, segments=[("RoomsPanel ", "b"),
    ("— room catalog + suggested-plan summary; cards have Select / quantity / remove controls that "
     "dispatch EDIT_ROOM_PLAN. Opening it triggers the room planner.", "")])
add_bullet(doc, segments=[("UnitDetailPanel ", "b"), ("— details for a unit the guest clicked in UE5.", "")])
add_bullet(doc, segments=[("DestinationsOverlay ", "b"), ("— hotel grid, shown in DESTINATION_SELECT (only Lake Como is active).", "")])
add_bullet(doc, segments=[("AmenitiesPanel ", "b"), ("— present but deprecated; amenity navigation is now voice-driven.", "")])
add_bullet(doc, segments=[("ChatPanel ", "b"),
    ("— text alternative to voice; swaps in for the avatar thumbnail using visibility:hidden (not "
     "unmount) to keep the HeyGen session alive.", "")])

# =========================================================================== #
# 10. LATENCY (current measured state)
# =========================================================================== #
add_h1(doc, "Latency — Current Measured State", 10)
add_para(doc, "An existing client report (reports/Voice Latency Performance Report.docx) analyzed ~108 "
    "production turns. Headline: typical AI-path response ~2.2 s, tail to 6 s; the LLM is ~80% of "
    "perceived delay; HeyGen adds a fixed ~700 ms TTS warmup; ~half of turns already use the sub-"
    "second fast path. The numbers below come from that analysis and the instrumentation in code.")
add_table(doc,
    ["Stage (AI-path turn)", "Typical", "Share"],
    [
        ["Client preparing request (debounce + classify + preflight)", "~315 ms", "~14%"],
        ["LLM call (OpenAI / Anthropic round-trip)", "~1,740 ms", "~79%"],
        ["Server-side processing (prompt build, validation)", "<10 ms", "<1%"],
        ["HeyGen voice-synthesis warmup", "~700 ms", "~32%"],
        ["End-to-end (overlapping, so shares exceed 100%)", "~2,200 ms", "—"],
    ],
    col_widths=[Cm(9.5), Cm(3.0), Cm(2.5)])

add_h2(doc, "10.1  Instrumentation you already have")
add_rich(doc, [("Everything is measured. ", "b"),
    ("lib/debug.ts defines a LatencyEntry with client walltimes (userSpeakEnded, userTranscription, "
     "intentClassified, orchestrateRequestSent/Received, repeatCalled, avatarSpeakStarted/Ended) and "
     "server timings (requestReceived, llmCallStart/End, responseSent, provider, model). "
     "orchestrate echoes serverTimings back so the client can split network vs server vs LLM. "
     "logLatency() prints [TURN-LATENCY] and POSTs to /api/log-latency, which writes per-session log "
     "files (newest 10 kept) to ./logs in dev or Vercel Blob in prod. Start every investigation here.", "")])

# =========================================================================== #
# 11. IMPROVEMENT BACKLOG
# =========================================================================== #
add_h1(doc, "Prioritized Improvement Backlog", 11)
add_para(doc, "Organized against the three mandates. Code entry points are given so work can start "
    "immediately.")

add_h2(doc, "11.1  Latency")
add_table(doc,
    ["Opportunity", "Why it matters", "Where", "Effort"],
    [
        ["Stream the LLM response", "No route streams today; speaking sentence-1 while the rest generates cuts 600-1000 ms off AI turns.", "orchestrate/route.ts + repeat() integration", "Med"],
        ["Expand fast-path coverage", "Many LLM turns (\"yes\", \"show me the pool\", \"next\") could be deterministic; each saves ~1.7 s.", "intents.ts, explorationDeterministic.ts, profileFastPath.ts", "Low-Med"],
        ["Tail-latency guard", "Abort LLM >3 s; fall back to a brief ack. Kills the 5-6 s worst case.", "useJourney.ts fetch timeouts", "Low"],
        ["Trim prompt / context", "Conversation history + tool schemas have grown; trim to ~8-10 turns and drop unusable tools.", "orchestrate/route.ts prompt builder", "Low"],
        ["Tune UE5 fade (3000 ms) + Vagon retry (1000 ms)", "Fixed per-navigation cost; make config-driven.", "ue5/bridge.ts, useUE5WebSocket.ts", "Low"],
        ["Append-only message persistence", "activeSession rewrites the whole array every 1.8 s.", "useIncrementalPersistence.ts", "Low"],
    ],
    col_widths=[Cm(4.6), Cm(6.2), Cm(4.3), Cm(1.3)])

add_h2(doc, "11.2  AI agent behaviour")
add_bullet(doc, segments=[("Profile-turn reliability: ", "b"),
    ("monitor validatorOverride rate in [ORCHESTRATE] logs; if it climbs, add few-shot examples or "
     "fall back to gpt-4o for PROFILE_COLLECTION (the code comment flags this exact path).", "")])
add_bullet(doc, segments=[("Ambiguous affirmatives in AMENITY_VIEWING: ", "b"),
    ("\"yes\" can mean \"tell me more\" or \"move on\". Ground intent in Ava's wording (\"Shall we "
     "move on to the spa?\" vs \"Want to hear more?\") so the next turn is deterministic.", "")])
add_bullet(doc, segments=[("Described-only amenities: ", "b"),
    ("the prompt forbids navigating to amenities without a UE5 scene, but enforcement is fuzzy; a "
     "dedicated tool with a constrained tool_choice would harden it.", "")])
add_bullet(doc, segments=[("Pre-hotel shortcut timing: ", "b"),
    ("the 3500 ms / 1.2 s two-phase travel is fragile; consider event-driven (wait for a UE5 "
     "scene-ready message) instead of fixed timers.", "")])

add_h2(doc, "11.3  Data extraction & analysis")
add_bullet(doc, segments=[("Guest composition & room allocation: ", "b"),
    ("regex is fragile on compound phrases (\"4 adults, 2 kids\", \"4 rooms for 8\"); these already "
     "defer to the LLM but the hand-offs are worth auditing in profileDeterministic.ts.", "")])
add_bullet(doc, segments=[("Honour the single-writer rule: ", "b"),
    ("the six profile-collection fields are orchestrate-owned; never let ProfileSync or new regex "
     "write them.", "")])
add_bullet(doc, segments=[("Persistence merges: ", "b"),
    ("personality/preferences merges are read-modify-write per profile slice (20-100 ms each); batch "
     "or move server-side to cut RTDB round-trips.", "")])
add_bullet(doc, segments=[("End-of-session analysis: ", "b"),
    ("/api/analyze-guest is best-effort with a 3 s timeout; consider skipping for very short sessions "
     "and enriching the guest-intelligence schema with structured objection outcomes.", "")])

# =========================================================================== #
# 12. ENVIRONMENT & COMMANDS
# =========================================================================== #
add_h1(doc, "Environment, Commands & Repo Map", 12)
add_h2(doc, "12.1  Environment variables")
add_table(doc,
    ["Variable", "Purpose"],
    [
        ["HEYGEN_API_KEY / HEYGEN_AVATAR_ID / HEYGEN_VOICE_ID / HEYGEN_CONTEXT_ID", "HeyGen LiveAvatar session + persona."],
        ["HEYGEN_API_URL (optional)", "Defaults to the LiveAvatar API host."],
        ["OPENAI_API_KEY", "All OpenAI routes (orchestrate non-profile, room-planner, intent, speech, POI, analyze-guest)."],
        ["ANTHROPIC_API_KEY", "Orchestrate during PROFILE_COLLECTION (claude-haiku-4-5)."],
        ["GOOGLE_PLACES_API_KEY (Places)", "POI verification/geolocation."],
        ["Firebase config keys", "Auth + Realtime DB + Storage."],
        ["NEXT_PUBLIC_STREAM_MODE", '"local" or "vagon" — selects UE5 transport.'],
        ["NEXT_PUBLIC_VAGON_STREAM_URL", "Local UE5 fallback (e.g. http://127.0.0.1)."],
    ],
    col_widths=[Cm(7.2), Cm(8.0)])
add_para(doc, "Confirm exact names against the API route source before deploying — some use "
    "NEXT_PUBLIC_-prefixed fallbacks server-side, which is misleading.", size=10, color=MUTED, italic=True)

add_h2(doc, "12.2  Commands")
add_bullet(doc, segments=[("npm run dev", "c"), ("  — start the Next dev server.", "")])
add_bullet(doc, segments=[("npx next build", "c"), ("  — production build (note: ", ""),
    ("next.config.mjs", "c"), (" sets ", ""), ("typescript.ignoreBuildErrors: true", "c"),
    (" and ", ""), ("images.unoptimized: true", "c"), (").", "")])
add_bullet(doc, segments=[("npx tsc --noEmit", "c"), ("  — type-check (ignore .next/dev/types noise from deleted pages).", "")])
add_bullet(doc, segments=[("python reports/_build_voice_latency_report.py", "c"),
    ("  — regenerate the latency client report (template for docx generation in this repo).", "")])

add_h2(doc, "12.3  Where to look first")
add_table(doc,
    ["Concern", "Start here"],
    [
        ["Turn orchestration / the hot loop", "lib/orchestrator/useJourney.ts"],
        ["The pure state machine", "lib/orchestrator/journey-machine.ts"],
        ["Server LLM decisions + prompts", "app/api/orchestrate/route.ts"],
        ["Intent regex + deterministic gates", "intents.ts, profileFastPath.ts, explorationDeterministic.ts"],
        ["Data extraction (client)", "lib/liveavatar/useUserProfile.ts, components/ProfileSync.tsx"],
        ["Behavioural analytics", "lib/guest-intelligence.ts, app/api/analyze-guest"],
        ["Avatar session / STT / TTS", "lib/liveavatar/context.tsx, useAvatarActions.ts, SandboxLiveAvatar.tsx"],
        ["UE5 protocol", "lib/ue5/bridge.ts, lib/useUE5WebSocket.ts"],
        ["State + persistence", "lib/omnam-store.tsx, lib/firebase/*"],
        ["Latency instrumentation", "lib/debug.ts, app/api/log-latency"],
        ["Page flow + boot", "app/home/page.tsx, app/layout.tsx"],
        ["Operator analytics UI", "app/admin/page.tsx"],
        ["Hotel/room/amenity data", "lib/hotel-data.ts, lib/hotels/lake-como.ts"],
    ],
    col_widths=[Cm(6.0), Cm(9.2)])

# ---- Closing note ---------------------------------------------------------- #
add_hr(doc)
add_para(doc,
    "Generated from a full read of the repository on the main branch (v0.4.3). Model names, "
    "temperatures, timeouts, and protocol strings in this document were verified directly against "
    "source. Where the checked-in CLAUDE.md and this report disagree, trust this report — it reflects "
    "the live code.", italic=True, color=MUTED, size=9.5)

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
