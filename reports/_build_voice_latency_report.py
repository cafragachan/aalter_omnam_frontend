"""
Generates the voice-latency client report as a .docx file.
Re-run this script to regenerate after editing the content below.

    python reports/_build_voice_latency_report.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT_PATH = Path(__file__).parent / "Voice Latency Performance Report.docx"

# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

ACCENT = RGBColor(0x1F, 0x3A, 0x68)   # dark blue
MUTED = RGBColor(0x55, 0x55, 0x55)
DIVIDER = RGBColor(0xCC, 0xCC, 0xCC)


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_table(doc, headers, rows, col_widths=None, header_fill="1F3A68"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    set_table_borders(table)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Body
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()
    return table


def add_para(doc, text, bold=False, italic=False, size=11, color=None, after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p


def add_rich_para(doc, segments, size=11, after=6):
    """segments = list of (text, bold) tuples."""
    p = doc.add_paragraph()
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

doc = Document()

# Page setup + base font
section = doc.sections[0]
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
trun = title.add_run("Voice Latency Performance Report")
trun.bold = True
trun.font.size = Pt(24)
trun.font.color.rgb = ACCENT

sub = doc.add_paragraph()
srun = sub.add_run("Omnam Concierge — AI Avatar Response Latency")
srun.font.size = Pt(13)
srun.font.color.rgb = MUTED
srun.italic = True

meta = doc.add_paragraph()
mrun = meta.add_run("Period analyzed: May 7 – May 11, 2026   ·   Sessions: 9 production sessions, ~108 user turns")
mrun.font.size = Pt(10)
mrun.font.color.rgb = MUTED

add_hr(doc)

# Executive Summary
add_h1(doc, "Executive Summary")
add_para(
    doc,
    "We instrumented the full voice pipeline — from the moment a guest stops speaking to the "
    "moment the avatar's voice begins — and measured every stage in between. After analyzing "
    "roughly 100 real production turns, the typical response time is around 2.2 seconds, with a "
    "long tail that can reach 6 seconds in worst cases. The dominant cost is the AI language "
    "model (Claude / OpenAI), which accounts for roughly 80% of perceived delay on turns that "
    "require reasoning. The system already handles the other half of turns through a \"fast path\" "
    "that bypasses the LLM and responds in under one second.",
)
add_para(
    doc,
    "There are clear, well-scoped opportunities to bring the typical response time below "
    "1.2 seconds without changing models or vendors.",
)

# What we measured
add_h1(doc, "What We Measured")
add_para(
    doc,
    "Each guest interaction passes through a sequence of stages, all of which we now log per turn "
    "into a per-session timing report:",
)
add_bullet(doc, "Speech recognition — HeyGen converts the guest's voice into text.")
add_bullet(doc, "Intent classification — our app determines what the guest wants.")
add_bullet(
    doc,
    "Response generation — for complex requests, an AI model (OpenAI or Anthropic) produces "
    "the avatar's reply; for simple commands, we respond directly.",
)
add_bullet(doc, "Voice synthesis — HeyGen converts the response text into the avatar's spoken audio.")

add_para(doc, "Two response paths exist:", after=4)
add_bullet(
    doc,
    "Fast path (~half of all turns): handled by our app directly. Typical end-to-end time: ~900 ms.",
)
add_bullet(
    doc,
    "AI path (~half of all turns): routed through the LLM. Typical end-to-end time: ~2,200 ms.",
)

# Findings
add_h1(doc, "Findings")
add_h2(doc, "Where the time goes (AI-path turns)")

add_table(
    doc,
    headers=["Stage", "Typical time", "Share of total"],
    rows=[
        ["Our app preparing the request", "~315 ms", "~14%"],
        ["AI model call", "~1,740 ms", "~79%"],
        ["Server-side processing", "<10 ms", "<1%"],
        ["Voice synthesis warmup (HeyGen)", "~700 ms", "~32%"],
        ["End-to-end total", "~2,200 ms", "—"],
    ],
)

add_para(
    doc,
    "(Stages overlap in the wall-clock total, so percentages don't sum to 100%.)",
    italic=True,
    size=10,
    color=MUTED,
)

add_h2(doc, "Key observations")

add_rich_para(doc, [
    ("1. The AI model is the dominant cost. ", True),
    ("On turns that require reasoning, the round-trip to OpenAI or Anthropic is by far the largest "
     "single contributor. The median is acceptable (~1.7 seconds) but the tail is a problem — we "
     "observed individual turns spiking to 3, 5, and even 5+ seconds.", False),
])
add_rich_para(doc, [
    ("2. Voice synthesis adds a fixed ~700 ms warmup. ", True),
    ("This is HeyGen's processing time between receiving the response text and playing the first "
     "audio frame. It is consistent across short and long responses, suggesting it's startup "
     "overhead rather than length-dependent rendering.", False),
])
add_rich_para(doc, [
    ("3. Half of all turns already use the fast path. ", True),
    ("These respond in ~900 ms and feel snappy. The other half go through the LLM — and many of "
     "those could be fast-path-eligible based on the guest's actual intent (e.g. \"yes please\", "
     "\"show me the rooms\", \"next one\"). Each conversion saves ~1.7 seconds per turn.", False),
])
add_rich_para(doc, [
    ("4. Tail latency is the worst part of the user experience. ", True),
    ("A 6-second response after a simple question creates the impression that \"the avatar is "
     "broken.\" These outliers are infrequent but disproportionately damage perceived quality.", False),
])
add_rich_para(doc, [
    ("5. Network and server-side processing are negligible. ", True),
    ("Our app's server adds essentially zero latency. Network hops to OpenAI/Anthropic are already "
     "part of the AI model call total. There is no infrastructure inefficiency to fix here.", False),
])

# Main issues
add_h1(doc, "Main Issues Observed")
add_table(
    doc,
    headers=["Issue", "Impact", "Frequency"],
    rows=[
        ["AI model call accounts for ~80% of response time",
         "High — sets the floor for typical latency",
         "Every AI-path turn"],
        ["Tail latency spikes to 5–6 seconds",
         "High — visibly breaks the experience",
         "~5% of turns"],
        ["Many turns route through the LLM unnecessarily",
         "Medium — doubles response time vs fast path",
         "~half of turns"],
        ["HeyGen voice synthesis warmup is fixed at ~700 ms",
         "Medium — sets a hard floor we cannot drop below today",
         "Every turn"],
    ],
)

# Proposed fixes
add_h1(doc, "Proposed Fixes")
add_para(doc, "Listed in priority order by expected impact.", italic=True, color=MUTED, after=10)

# Fix 1
add_h2(doc, "1. Stream the AI response and start speaking earlier")
add_rich_para(doc, [("What: ", True), (
    "Today, we wait for the full AI response to finish generating before sending it to the avatar. "
    "Streaming would let us send the response sentence-by-sentence — the avatar can start speaking "
    "the first sentence while the rest is still being generated.", False)])
add_rich_para(doc, [("Expected impact: ", True), (
    "Cuts the perceived delay on AI-path turns by 600–1,000 ms. A 5-second tail latency would "
    "feel like ~2 seconds.", False)])
add_rich_para(doc, [("Effort: ", True), (
    "Medium. Requires coordinated changes between our orchestration code and the HeyGen integration.",
    False)])

# Fix 2
add_h2(doc, "2. Expand fast-path coverage")
add_rich_para(doc, [("What: ", True), (
    "Audit which guest utterances currently trigger the AI path and identify ones that could be "
    "handled deterministically. Examples: \"yes\", \"no\", \"show me the pool\", \"cheaper option\".",
    False)])
add_rich_para(doc, [("Expected impact: ", True), (
    "Each converted turn drops from ~2.2 seconds to ~0.9 seconds. If we can move 30–40% of "
    "current AI turns to fast path, the session-average response time drops by ~40%.", False)])
add_rich_para(doc, [("Effort: ", True), (
    "Low to medium. Mostly intent-pattern work; no model changes.", False)])

# Fix 3
add_h2(doc, "3. Tail latency mitigation")
add_rich_para(doc, [("What: ", True), (
    "When an AI call exceeds ~3 seconds, abort and either retry, or fall back to a brief "
    "acknowledgement (\"one moment…\") while the slow call continues.", False)])
add_rich_para(doc, [("Expected impact: ", True), (
    "Eliminates the 5-6 second worst-case experience entirely. Median is unaffected.", False)])
add_rich_para(doc, [("Effort: ", True), ("Low.", False)])

# Fix 4
add_h2(doc, "4. Prompt and context optimization")
add_rich_para(doc, [("What: ", True), (
    "The conversation history and tool descriptions sent with each LLM call have grown over "
    "time. Trimming to the last 8-10 turns and removing tool schemas the model can't use in the "
    "current stage shrinks each request.", False)])
add_rich_para(doc, [("Expected impact: ", True), (
    "Modest — typically 200–400 ms off the AI model call.", False)])
add_rich_para(doc, [("Effort: ", True), ("Low.", False)])

# Expected outcomes
add_h1(doc, "Expected Outcomes")
add_para(doc, "If we implement fixes 1, 2, and 3:")
add_table(
    doc,
    headers=["Metric", "Today", "Projected"],
    rows=[
        ["Typical AI-path response", "2.2 s", "~1.2 s"],
        ["Worst-case response", "6+ s", "<3 s"],
        ["Fast-path coverage", "~50% of turns", "~75% of turns"],
        ["Session-average perceived latency", "1.6 s", "<1 s"],
    ],
)
add_para(
    doc,
    "This would move the experience from \"noticeably AI-mediated\" to \"conversational.\"",
    italic=True,
)

# Next steps
add_h1(doc, "Next Steps")
add_para(doc, "We recommend proceeding in this order:")
add_bullet(doc, "Fix #3 (tail latency mitigation) — fastest to ship, biggest reduction in worst-case pain.")
add_bullet(doc, "Fix #2 (fast-path expansion) — biggest impact on the typical session.")
add_bullet(doc, "Fix #1 (streaming) — largest single perceived-latency win, longer to implement.")
add_bullet(doc, "Fix #4 (prompt trimming) — incremental, can be done alongside the above.")
add_para(
    doc,
    "Each fix is independent and ships value on its own. We can begin work on any of them upon approval.",
)

# Save
doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
